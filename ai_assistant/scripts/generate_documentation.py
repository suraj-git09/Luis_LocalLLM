"""Generate comprehensive project documentation as a Word file.

This script produces AI_Assistant_Project_Documentation.docx.
Run it after major changes:

    python scripts/generate_documentation.py

The generated document is a snapshot. The primary living documentation
is in README.md (root), AIAssistantAndroid/README.md, SECURITY.md, and .env.example.
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parent.parent / "AI_Assistant_Project_Documentation.docx"


def set_document_styles(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    for level, size in [(1, 18), (2, 14), (3, 12)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Arial"
        h.font.bold = True
        h.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        h.font.size = Pt(size)


def add_toc(paragraph):
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    run._r.append(fld)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()


def add_page_numbers(doc: Document):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Luis AI Assistant Documentation | Page ")
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        run._r.append(fld)


def build_document() -> Document:
    doc = Document()
    set_document_styles(doc)

    # --- COVER PAGE ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Luis — AI Assistant\n")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run("Professional Project Documentation\nWeb UI • Android Client • Authentication • History")
    s.font.size = Pt(15)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"\n\nVersion 1.1\n{date.today().strftime('%B %d, %Y')}\n")
    meta.add_run("\nAuthor: Suraj Raj\n")
    meta.add_run("Key updates: Email OTP, Conversation History, New Chat, Google Sign-In removed\n")

    doc.add_page_break()

    # --- TABLE OF CONTENTS ---
    doc.add_heading("Table of Contents", level=1)
    toc = doc.add_paragraph()
    add_toc(toc)
    doc.add_page_break()

    # --- 1. INTRODUCTION ---
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "Imagine you have a very smart helper on your computer. You type a question or speak to it, "
        "and it tries to answer you. Sometimes it uses built-in tools (like a calculator). "
        "Sometimes it looks in its memory. Sometimes it asks a big AI brain (LLM) for help."
    )
    doc.add_paragraph(
        "This project is called AI Assistant. It is a production-ready Python application that can run "
        "in three ways: Text Mode (keyboard), Voice Mode (microphone + speaker), and Web Mode (browser UI)."
    )

    doc.add_heading("1.1 What Problem Does It Solve?", level=2)
    add_bullet(doc, "People want one assistant that works offline when possible.")
    add_bullet(doc, "People want safe local commands (notes, reminders, math) without always paying for cloud AI.")
    add_bullet(doc, "People want smart answers for open questions when local tools are not enough.")
    add_bullet(doc, "Developers want clean, modular code that is easy to extend.")

    doc.add_heading("1.2 Simple Analogy for Beginners", level=2)
    doc.add_paragraph(
        "Think of the assistant like a school reception desk:\n"
        "1) You walk in and say what you need (your message).\n"
        "2) A receptionist figures out what type of request it is (intent classifier).\n"
        "3) If it is a simple task, a specific worker handles it (command).\n"
        "4) If it is a hard question, the receptionist asks the expert teacher (LLM).\n"
        "5) The answer comes back to you."
    )

    # --- 2. TECH STACK ---
    doc.add_heading("2. Technology Stack (Explained Simply)", level=1)
    doc.add_paragraph("Each technology below is a building block. Together they make the assistant work.")

    tech_rows = [
        ("Python 3.11+", "Main programming language", "Like the language we write instructions in. Easy to read and very popular for AI apps."),
        ("asyncio", "Async runtime", "Lets the app do waiting tasks (like LLM calls) without freezing everything."),
        ("python-dotenv", "Environment loader", "Reads secret settings from .env file so keys are not hard-coded in code."),
        ("OpenAI Python SDK", "LLM API client", "Talks to AI models using the same style as OpenAI API. Works with Ollama too."),
        ("Ollama", "Local LLM server", "Runs AI models on your own computer. No internet needed for AI answers."),
        ("Flask", "Web framework", "Creates the browser chat website at http://127.0.0.1:5000."),
        ("SQLite", "Embedded database", "Tiny database stored in files. Used for cache and reminders."),
        ("JSON files", "Simple storage", "Stores notes in a human-readable list file."),
        ("Vosk", "Wake-word STT", "Listens for the word 'assistant' using offline speech model."),
        ("Faster-Whisper", "Speech-to-text", "Converts your voice into text. Good with accents."),
        ("pyttsx3", "Text-to-speech", "Speaks answers aloud using Windows voice engine."),
        ("pytest", "Testing framework", "Automatically checks that features still work after changes."),
        ("Docker", "Container packaging", "Runs the app in a portable box with Ollama sidecar."),
        ("GitHub Actions", "CI pipeline", "Runs tests automatically when code is pushed."),
    ]
    add_table(doc, ["Technology", "Role", "Simple Explanation"], tech_rows)

    # --- 3. ARCHITECTURE ---
    doc.add_heading("3. System Architecture", level=1)
    doc.add_paragraph(
        "The project follows a modular architecture. Each folder has one job. "
        "This is like organizing a kitchen: knives in one drawer, plates in another."
    )

    arch_rows = [
        ("main.py", "Front door", "Starts the app and lets user pick text, voice, or web mode."),
        ("app/", "Factory wiring", "Connects all parts together (bootstrap)."),
        ("assistant/", "Brain pipeline", "Receives input and routes it to the right handler."),
        ("commands/", "Tool workers", "Each command does one job: math, notes, weather, etc."),
        ("nlp/", "Intent guesser", "Figures out what the user wants from their words."),
        ("services/", "Support systems", "LLM, cache, logging, voice, reminders, health."),
        ("config/", "Settings", "Reads .env and provides configuration values."),
        ("utils/", "Safety helpers", "Validates and cleans user input."),
        ("web/", "Browser UI", "Flask server + HTML/CSS/JS chat interface."),
        ("tests/", "Quality checks", "Automated tests for important features."),
        ("data/", "Runtime storage", "Databases, logs, notes, optional voice models."),
        ("scripts/", "Helper launchers", "Run web, health checks, production startup."),
    ]
    add_table(doc, ["Folder/File", "Role Name", "What It Does"], arch_rows)

    doc.add_heading("3.1 Request Flow (Step by Step)", level=2)
    flow_steps = [
        "User sends a message (keyboard, voice, or web chat).",
        "main.py or web/server.py receives it.",
        "AssistantCore validates input length and blocks bad characters.",
        "IntentClassifier predicts intent (help, notes, calculate, general_qa, etc.).",
        "Router picks the matching Command from CommandRegistry.",
        "Command runs. If it returns None, GeneralQACommand takes over.",
        "GeneralQA checks cache, then offline knowledge base, then LLM.",
        "Prompt Optimizer picks token budget and system style for LLM.",
        "LLMService calls Ollama or online API and returns answer.",
        "Response is shown to user. Metrics and logs are updated.",
    ]
    for i, step in enumerate(flow_steps, 1):
        add_bullet(doc, f"Step {i}: {step}")

    # --- 4. FOLDER DOCUMENTATION ---
    doc.add_heading("4. Folder-by-Folder Documentation", level=1)

    folders = {
        "app/": {
            "purpose": "Application bootstrap layer. Wires all services and commands together.",
            "why": "Instead of messy setup code in main.py, we use one clean factory function.",
            "files": {
                "__init__.py": "Marks app as a Python package and exports build_application.",
                "bootstrap.py": "build_application() creates cache, storage, LLM, router, assistant, worker, health service.",
            },
        },
        "assistant/": {
            "purpose": "Core assistant pipeline: input handling and routing.",
            "why": "Separates 'how we process a message' from individual command logic.",
            "files": {
                "__init__.py": "Package marker for assistant module.",
                "core.py": "AssistantCore.handle_input() validates input and calls router.",
                "router.py": "Router.route() picks command by intent and handles errors.",
                "fallback.py": "OfflineFallback gives cached/default message when command fails.",
            },
        },
        "commands/": {
            "purpose": "Pluggable command handlers. Each command is one feature.",
            "why": "Easy to add new features without breaking existing ones.",
            "files": {
                "__init__.py": "Package marker.",
                "base.py": "Abstract Command class with can_handle() and execute().",
                "registry.py": "CommandRegistry stores and finds commands by intent name.",
                "calculator.py": "Safe math using AST parser (no dangerous eval).",
                "notes.py": "save note / show notes using JSON storage.",
                "system_info.py": "Shows OS info and current time.",
                "weather.py": "Offline sample weather for Delhi, Mumbai, London.",
                "reminders.py": "Parses 'remind me in X minutes' and stores reminders.",
                "help.py": "Lists supported commands and examples.",
                "health.py": "Returns system health and metrics report.",
                "general_qa.py": "Cache -> knowledge base -> LLM fallback chain.",
            },
        },
        "config/": {
            "purpose": "Central configuration from environment variables.",
            "why": "One place for all settings; easy to change without editing code.",
            "files": {
                "__init__.py": "Exports Settings and get_settings.",
                "settings.py": "Settings dataclass + get_settings() reads .env values.",
            },
        },
        "nlp/": {
            "purpose": "Natural Language Processing - intent detection.",
            "why": "Maps free-form user text to structured command names.",
            "files": {
                "__init__.py": "Package marker.",
                "intent_classifier.py": "Keyword-based intent classifier (help, notes, calculate, etc.).",
            },
        },
        "services/": {
            "purpose": "Reusable backend services used by commands and bootstrap.",
            "why": "Shared logic (LLM, cache, logging) should not be duplicated.",
            "files": {
                "__init__.py": "Package marker.",
                "cache.py": "SQLite cache for question-answer pairs.",
                "storage.py": "JSON notes storage with corruption recovery.",
                "conversation_memory.py": "Keeps recent chat messages for LLM context.",
                "knowledge_base.py": "Small offline Q&A dictionary (Einstein, Python, AI, etc.).",
                "llm_service.py": "Connects to Ollama or online API; token limits and timeouts.",
                "prompt_optimizer.py": "Chooses response style and token budget per prompt.",
                "rate_limiter.py": "Limits how many LLM calls per minute.",
                "reminder_store.py": "SQLite persistence for scheduled reminders.",
                "scheduler.py": "Background worker thread that fires due reminders.",
                "logging_config.py": "Console + rotating file logs.",
                "metrics.py": "Counts requests, errors, reminders delivered.",
                "health.py": "Checks DB, logs, worker, LLM; formats health report.",
                "voice_stt.py": "Offline wake word (Vosk) + Whisper transcription.",
                "voice_tts.py": "Offline text-to-speech using pyttsx3.",
            },
        },
        "utils/": {
            "purpose": "Cross-cutting utility helpers.",
            "why": "Security validation should be reusable across commands.",
            "files": {
                "__init__.py": "Exports validation helpers.",
                "validation.py": "sanitize_text(), length limits, control character filtering.",
            },
        },
        "web/": {
            "purpose": "Flask web application for browser-based chat UI.",
            "why": "Modern visual interface for users who prefer browser over terminal.",
            "files": {
                "__init__.py": "Exports create_web_app and run_web_server.",
                "server.py": "Flask routes: /, /api/chat, /api/health.",
                "templates/index.html": "Main chat page layout with sidebar and composer.",
                "static/css/app.css": "Dark theme styling, responsive layout.",
                "static/js/app.js": "Chat logic, API calls, code formatting, health polling.",
            },
        },
        "tests/": {
            "purpose": "Automated test suite using pytest.",
            "why": "Catch bugs early and prove features work.",
            "files": {
                "conftest.py": "Adds project root to Python path for tests.",
                "test_calculator.py": "Tests math and blocks unsafe expressions.",
                "test_intent.py": "Tests intent classification.",
                "test_reminders.py": "Tests reminder create/list/due detection.",
                "test_storage.py": "Tests corrupt notes file recovery.",
                "test_security.py": "Tests input validation and memory trimming.",
                "test_health.py": "Tests health report generation.",
                "test_llm_service.py": "Tests Ollama model resolution fallback.",
                "test_prompt_optimizer.py": "Tests token profile selection.",
            },
        },
        "scripts/": {
            "purpose": "Operational helper scripts.",
            "why": "One-command startup for users and deployment tools.",
            "files": {
                "run.ps1": "Windows production run: venv, install, test, text mode.",
                "run.sh": "Linux/macOS production run script.",
                "run_web.ps1": "Starts Flask web UI on port 5000.",
                "healthcheck.py": "Exit 0 if healthy; used by Docker HEALTHCHECK.",
                "generate_documentation.py": "This script that created this Word document.",
            },
        },
        "data/": {
            "purpose": "Runtime data directory (created automatically).",
            "why": "Keeps databases, logs, and models separate from source code.",
            "files": {
                ".gitkeep": "Placeholder so Git tracks empty data folder.",
                "cache.db": "SQLite cache (auto-created at runtime).",
                "notes.json": "Saved notes list (auto-created).",
                "reminders.db": "Scheduled reminders database.",
                "logs/assistant.log": "Rotating application log file.",
                "vosk-model-small-en-us-0.15/": "Optional offline wake-word model (downloaded separately).",
            },
        },
    }

    for folder, info in folders.items():
        doc.add_heading(f"4.{list(folders.keys()).index(folder)+1} {folder}", level=2)
        doc.add_paragraph(f"Purpose: {info['purpose']}")
        doc.add_paragraph(f"Why used here: {info['why']}")
        file_rows = [(fname, desc) for fname, desc in info["files"].items()]
        add_table(doc, ["File", "What It Does (Simple Words)"], file_rows)

    # --- 5. ROOT FILES ---
    doc.add_heading("5. Root-Level Files", level=1)
    root_files = [
        ("main.py", "Main entry point. Supports --mode text|voice|web. Starts assistant."),
        (".env.example", "Template for environment variables. Copy to .env and fill values."),
        (".env", "Your private local config (never commit to Git)."),
        (".gitignore", "Tells Git which files to ignore (secrets, cache, venv)."),
        (".dockerignore", "Tells Docker which files to skip when building image."),
        ("requirements.txt", "Full Python dependencies (core + voice + test)."),
        ("requirements-prod.txt", "Minimal dependencies for production Docker image."),
        ("pyproject.toml", "Project metadata, optional dependency groups, pytest config."),
        ("README.md", "Quick start guide and feature overview."),
        ("SECURITY.md", "Security policy, secret handling, key rotation checklist."),
        ("Dockerfile", "Instructions to build production container image."),
        ("docker-compose.yml", "Runs assistant + Ollama together with volumes."),
    ]
    add_table(doc, ["File", "Explanation"], root_files)

    # --- 6. COMMANDS ---
    doc.add_heading("6. Supported Commands", level=1)
    cmd_rows = [
        ("help", "help", "Shows what the assistant can do."),
        ("calculate", "calculate 12 + 8", "Does math safely."),
        ("notes", "save note buy milk", "Saves a note to JSON file."),
        ("notes", "show notes", "Lists all saved notes."),
        ("system_info", "system info", "Shows computer OS and processor info."),
        ("system_info", "what is the time", "Shows current date and time."),
        ("weather", "weather in delhi", "Returns offline sample weather."),
        ("reminder", "remind me to stretch in 5 minutes", "Creates timed reminder."),
        ("reminder", "show reminders", "Lists pending reminders."),
        ("health", "system health", "Shows component status and metrics."),
        ("general_qa", "any open question", "Uses cache, knowledge, then LLM."),
    ]
    add_table(doc, ["Intent", "Example", "What Happens"], cmd_rows)

    # --- 7. ENV VARS ---
    doc.add_heading("7. Environment Variables Reference", level=1)
    env_rows = [
        ("OPENAI_API_KEY", "—", "Online AI key when Ollama is not used."),
        ("OPENAI_BASE_URL", "api.openai.com/v1", "URL of online LLM service."),
        ("OPENAI_MODEL", "gpt-4o-mini", "Online model name."),
        ("LOCAL_MODEL", "gemma2:2b", "Ollama model name (must be installed)."),
        ("OLLAMA_BASE_URL", "localhost:11434/v1", "Local Ollama server address."),
        ("LOG_LEVEL", "INFO", "How much detail in logs."),
        ("DATA_DIR", "data", "Where databases and logs are stored."),
        ("VOSK_MODEL_DIR", "data/vosk-model...", "Path to wake-word speech model."),
        ("WHISPER_MODEL_SIZE", "tiny", "Voice transcription model size."),
        ("WAKE_WORD", "assistant", "Word that activates voice mode."),
        ("TTS_RATE", "175", "Speech speed for voice output."),
        ("TTS_VOLUME", "1.0", "Speech volume."),
        ("MAX_INPUT_CHARS", "2000", "Maximum user message length."),
        ("MAX_NOTE_CHARS", "500", "Maximum note length."),
        ("MAX_REMINDER_CHARS", "300", "Maximum reminder text length."),
        ("MAX_MEMORY_MESSAGES", "8", "How many past messages LLM remembers."),
        ("LLM_TIMEOUT_SECONDS", "30", "Max wait time for AI response."),
        ("LLM_MAX_TOKENS", "256", "Upper cap on AI response length."),
        ("LLM_RATE_LIMIT_CALLS", "10", "Max LLM calls per period."),
        ("LLM_RATE_LIMIT_PERIOD", "60", "Rate limit window in seconds."),
        ("LOG_MAX_BYTES", "5242880", "Log file size before rotation (5 MB)."),
        ("LOG_BACKUP_COUNT", "3", "How many old log files to keep."),
    ]
    add_table(doc, ["Variable", "Default", "Meaning"], env_rows)

    # --- 8. MODES ---
    doc.add_heading("8. Running Modes", level=1)
    doc.add_heading("8.1 Text Mode", level=2)
    doc.add_paragraph("Command: python main.py --mode text\nBest for: Terminal users, debugging, lightweight usage.")
    doc.add_heading("8.2 Voice Mode", level=2)
    doc.add_paragraph(
        "Command: python main.py --mode voice\n"
        "Best for: Hands-free usage. Say 'assistant' wake word, then speak your request."
    )
    doc.add_heading("8.3 Web Mode", level=2)
    doc.add_paragraph(
        "Command: python main.py --mode web\n"
        "Open: http://127.0.0.1:5000\n"
        "Best for: Visual chat UI with quick actions, health panel, and code formatting."
    )

    # --- 9. TOKEN OPTIMIZATION ---
    doc.add_heading("9. Token Optimization (Smart AI Usage)", level=1)
    doc.add_paragraph(
        "Tokens are like word coins for AI. Fewer tokens = faster and cheaper responses. "
        "The Prompt Optimizer studies your message and picks the right response style:"
    )
    token_rows = [
        ("code", "280-380 tokens", "Code requests: returns mostly code, little explanation."),
        ("brief", "64 tokens", "When you say 'briefly': one short sentence."),
        ("concise", "96 tokens", "Short questions: 1-2 sentences."),
        ("list", "160 tokens", "Bullet list, max 5 points."),
        ("explain", "200 tokens", "Clear short explanation."),
        ("balanced/detailed", "140-220 tokens", "Normal questions with matched length."),
    ]
    add_table(doc, ["Style", "Token Budget", "When Used"], token_rows)

    # --- 10. SECURITY ---
    doc.add_heading("10. Security Features", level=1)
    sec_points = [
        "No hardcoded API keys — all secrets come from .env.",
        "Safe calculator uses AST parsing, not dangerous eval().",
        "Input validation blocks oversized text and control characters.",
        "SQLite uses parameterized queries to prevent injection.",
        "Rate limiting prevents LLM abuse.",
        "Log rotation prevents disk filling.",
        ".gitignore prevents committing secrets and runtime data.",
    ]
    for p in sec_points:
        add_bullet(doc, p)

    # --- 11. DEPLOYMENT ---
    doc.add_heading("11. Deployment Options", level=1)
    deploy_points = [
        "Local: scripts/run.ps1 or scripts/run.sh",
        "Web: scripts/run_web.ps1 or python main.py --mode web",
        "Docker: docker compose up --build",
        "CI: GitHub Actions runs pytest + healthcheck on every push",
    ]
    for p in deploy_points:
        add_bullet(doc, p)

    # --- 12. TESTING ---
    doc.add_heading("12. Testing", level=1)
    doc.add_paragraph("Run all tests: pytest\nCurrent suite: 21 tests covering calculator, intent, reminders, storage, security, health, LLM, and prompt optimizer.")

    # --- 13. TROUBLESHOOTING ---
    doc.add_heading("13. Troubleshooting Guide", level=1)
    trouble_rows = [
        ("Offline answer / AI unavailable", "Start Ollama or set OPENAI_API_KEY in .env."),
        ("Model not found", "Run ollama list and set LOCAL_MODEL to an installed model."),
        ("Voice not working", "Install voice deps and download Vosk model."),
        ("Web page not loading", "Run python main.py --mode web and open port 5000."),
        ("Notes error", "Corrupt notes.json is auto-reset to empty list."),
    ]
    add_table(doc, ["Problem", "Solution"], trouble_rows)

    # --- 14. GLOSSARY ---
    doc.add_heading("14. Glossary (Kid-Friendly)", level=1)
    glossary = [
        ("API", "A way for programs to talk to other programs over the internet."),
        ("Intent", "What the user is trying to do (help, calculate, ask question)."),
        ("LLM", "Large Language Model — a very big AI trained on lots of text."),
        ("Ollama", "Software that runs AI models on your own computer."),
        ("Token", "A small piece of text the AI reads or writes. Like word coins."),
        ("Cache", "Saved answer memory so repeated questions are faster."),
        ("Router", "Traffic director that sends each message to the right handler."),
        ("Bootstrap", "Startup wiring that connects all parts when app launches."),
        ("Flask", "A simple Python tool for making websites."),
        ("SQLite", "A tiny database stored in one file."),
    ]
    add_table(doc, ["Term", "Simple Meaning"], glossary)

    # --- 15. CONCLUSION ---
    doc.add_heading("15. Conclusion", level=1)
    doc.add_paragraph(
        "The AI Assistant project is a complete, production-style learning and utility platform. "
        "It teaches modular architecture, offline-first design, AI integration, security, monitoring, "
        "testing, and deployment. Every folder and file has a clear responsibility, making the project "
        "easy to understand, extend, and present in portfolios or industrial demonstrations."
    )

    doc.add_paragraph("\n— End of Document —").alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_numbers(doc)
    return doc


if __name__ == "__main__":
    document = build_document()
    document.save(OUTPUT)
    print(f"Documentation saved to: {OUTPUT}")